#!/usr/bin/env python3
"""Build Waseem Ansari's CV (.docx) matching the portfolio site."""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ACCENT = RGBColor(0x2F, 0x4C, 0xC4)      # deep periwinkle, from the portfolio
INK = RGBColor(0x14, 0x18, 0x24)         # near-black
MUTED = RGBColor(0x54, 0x5C, 0x70)       # cool grey
RULE = "C9CDDA"
BODY_FONT = "Calibri"

doc = Document()

# ---------- page setup (A4) ----------
sec = doc.sections[0]
sec.page_width = Inches(8.27)
sec.page_height = Inches(11.69)
for attr, val in (("left_margin", 0.6), ("right_margin", 0.6),
                  ("top_margin", 0.5), ("bottom_margin", 0.5)):
    setattr(sec, attr, Inches(val))

# ---------- base style ----------
normal = doc.styles["Normal"]
normal.font.name = BODY_FONT
normal.font.size = Pt(9.5)
normal.font.color.rgb = INK
normal._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
pf = normal.paragraph_format
pf.space_after = Pt(0)
pf.space_before = Pt(0)
pf.line_spacing = 1.12


def para(space_before=0, space_after=0, line=1.12):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = line
    return p


def run(p, text, size=9.5, bold=False, italic=False, color=INK,
        caps=False, spacing=None, font=BODY_FONT):
    r = p.add_run(text)
    r.font.name = font
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    r.font.color.rgb = color
    r._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    if caps:
        r.font.all_caps = True
    if spacing is not None:  # letter-spacing in twentieths of a point
        sp = OxmlElement("w:spacing")
        sp.set(qn("w:val"), str(spacing))
        r._element.get_or_add_rPr().append(sp)
    return r


def bottom_rule(p, color=RULE, size=6):
    pPr = p._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "3")
    bottom.set(qn("w:color"), color)
    borders.append(bottom)
    pPr.append(borders)


def heading(text, first=False):
    p = para(space_before=0 if first else 11, space_after=5)
    run(p, text, size=10, bold=True, color=ACCENT, caps=True, spacing=28)
    bottom_rule(p)
    return p


def role(title, org, dates, subtitle=None):
    """Job title line with right-aligned dates, then the organisation."""
    p = para(space_before=7, space_after=1)
    tabs = p.paragraph_format.tab_stops
    tabs.add_tab_stop(Inches(7.07), WD_TAB_ALIGNMENT.RIGHT)
    run(p, title, size=10.5, bold=True)
    run(p, "\t")
    run(p, dates, size=9, bold=True, color=MUTED)
    q = para(space_after=3)
    run(q, org, size=9.5, italic=True, color=ACCENT)
    if subtitle:
        run(q, "  |  " + subtitle, size=9, italic=True, color=MUTED)


def bullet(text, space_after=1.5):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.19)
    p.paragraph_format.first_line_indent = Inches(-0.15)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.1
    run(p, text, size=9.5)
    return p


# =====================================================================
# HEADER
# =====================================================================
p = para(space_after=1)
run(p, "WASEEM AHMAD ANSARI", size=22, bold=True, spacing=12)

p = para(space_after=4)
run(p, "Senior Software Engineer  ·  AI & LLM Quality Evaluator  ·  Search Evaluator",
    size=10.5, color=ACCENT, bold=True)

p = para(space_after=2)
run(p, "+91 92703 04741   |   Waseemwdd0165@gmail.com   |   "
       "linkedin.com/in/waseem-ahmad-ansari-bba5771ab   |   waseemauto.com   |   Mumbai, India",
    size=8.2, color=MUTED)
bottom_rule(p, color="2F4CC4", size=10)

# =====================================================================
# PROFILE
# =====================================================================
heading("Profile")
p = para(space_after=3)
run(p, "Computer Engineer (B.E.) with nine years as a senior software engineer building software "
       "applications, websites, and mobile apps, and seven years evaluating the search and AI systems "
       "that shape what people find online. The two sides feed each other: knowing how a search stack "
       "or a language model pipeline is put together makes for sharper judgement of what it produces, "
       "and years of assessing that output against detailed guidelines make for more careful engineering. "
       "Comfortable across front-end and back-end development, functional and non-functional QA, "
       "LLM evaluation and alignment work, prompt engineering, and multilingual localization. "
       "Based in Mumbai and open to remote or on-site work.")

# =====================================================================
# CORE COMPETENCIES
# =====================================================================
heading("Core Competencies")

skills = [
    ("Software Engineering & QA",
     "Software engineering · Desktop applications · Application implementation · Functional & "
     "non-functional testing · UI development · Agile delivery · Code standards & mentoring"),
    ("Web & Mobile Development",
     "Responsive web design · Dynamic & static websites · Mobile app development · "
     "Cross-platform apps · Front-end & back-end integration"),
    ("LLM Evaluation & Alignment",
     "RLHF & response ranking · Preference data · Red teaming · Safety & policy review · "
     "Hallucination detection · Factuality checking · Rubric application · SFT data authoring"),
    ("AI, LLM & Prompt Engineering",
     "Prompt engineering · LLM optimization · Generative AI quality analysis · Model comparison · "
     "Chatbot & voice assistant testing · AI model training · AI content review"),
    ("AI Data Operations",
     "AI data quality assurance · ML data operations · Data validation · Data annotation · "
     "Multimodal annotation · Image, video & audio labelling · Dataset curation"),
    ("Search & Ads Evaluation",
     "Search engine algorithm evaluation · Information retrieval analysis · "
     "Digital advertisement quality · Relevance rating"),
    ("Localization & Linguistics",
     "Multilingual AI evaluation · Linguistic data · Transcription · Translation · "
     "Audio localization · Dubbing · Lip sync"),
]

tbl = doc.add_table(rows=0, cols=2)
tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
tbl.autofit = False
col_w = [Inches(1.72), Inches(5.35)]
tbl.columns[0].width = col_w[0]
tbl.columns[1].width = col_w[1]

for k, v in skills:
    row = tbl.add_row()
    for i, cell in enumerate(row.cells):
        cell.width = col_w[i]
        cp = cell.paragraphs[0]
        cp.paragraph_format.space_before = Pt(1.5)
        cp.paragraph_format.space_after = Pt(1.5)
        cp.paragraph_format.line_spacing = 1.1
    run(row.cells[0].paragraphs[0], k, size=9.2, bold=True)
    run(row.cells[1].paragraphs[0], v, size=9.2, color=MUTED)

# remove table borders and left padding
tblPr = tbl._tbl.tblPr
borders = OxmlElement("w:tblBorders")
for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
    e = OxmlElement(f"w:{edge}")
    e.set(qn("w:val"), "none")
    borders.append(e)
tblPr.append(borders)
ind = OxmlElement("w:tblInd")
ind.set(qn("w:w"), "0")
ind.set(qn("w:type"), "dxa")
tblPr.append(ind)

# =====================================================================
# EXPERIENCE
# =====================================================================
heading("Professional Experience")

role("Senior Software Engineer",
     "Net Tech Services India Pvt. Ltd, Mumbai",
     "Jan 2017 – Mar 2025")
for b in [
    "Designed, built, and shipped software applications, websites, and mobile apps for clients, "
    "from first architecture decisions through to deployment.",
    "Led front-end and back-end integration work, keeping data flowing cleanly and the system secure.",
    "Ran thorough functional and non-functional testing, which meaningfully cut the number of bugs "
    "that reached production.",
    "Worked closely with UI/UX designers, product managers, and QA in an Agile setup, usually against "
    "tight deadlines.",
    "Mentored junior developers and helped settle on coding standards the whole team could work to.",
]:
    bullet(b)

role("Support Executive Officer",
     "Net Tech Services India Pvt. Ltd, Mumbai",
     "2016 – 2016")
for b in [
    "Tested core software products, both functionally and for performance and reliability.",
    "Supported application rollouts and pitched in on UI development.",
]:
    bullet(b)

role("AI & LLM Quality Evaluator  ·  Search Evaluator  ·  Localization Analyst",
     "Freelance & contract work",
     "2018 – Present")
for b in [
    "Rank competing model responses side by side to produce the preference data behind RLHF, "
    "applying detailed written rubrics consistently across high volumes of work.",
    "Red-team models for unsafe, biased, or policy-violating responses, and document what surfaces "
    "so it can be addressed.",
    "Check model claims against sources to catch hallucinated facts and fabricated citations.",
    "Write and refine prompts, and author ideal instruction–response pairs used in supervised fine-tuning.",
    "Evaluate model output across multiple languages, drawing on years of localization and linguistic work.",
    "Test chatbots and voice assistants for accuracy, tone, and whether they actually complete the task "
    "asked of them.",
    "Annotate and evaluate multimodal data — image, video, and audio alongside text — and validate "
    "training datasets before they reach a model.",
    "Rate search results and digital ads against quality guidelines, helping improve how accurately "
    "systems rank and match what people are looking for.",
    "Handle linguistic data work — transcription, translation, audio localization, dubbing, and lip sync — "
    "and design and build responsive websites for clients.",
]:
    bullet(b)

# =====================================================================
# EDUCATION
# =====================================================================
heading("Education")

role("Bachelor of Engineering (B.E.), Computer Engineering",
     "University of Mumbai — Rizvi Education Society's College of Engineering",
     "2010 – 2014")
p = para(space_after=2)
run(p, "Examination held May 2015 · Second Class · Convocation 16 January 2017 · "
       "Certificate no. 0032270", size=9, color=MUTED)

role("Higher Secondary School Certificate (HSC)",
     "The Malegaon Jr. College",
     "2007 – 2009")

# =====================================================================
# ADDITIONAL INFORMATION
# =====================================================================
heading("Additional Information")

extras = [
    ("Availability",
     "Open to full-time, contract, and freelance engagements. Available for remote work "
     "or on-site in Mumbai."),
    ("Working Style",
     "Nine years delivering to client deadlines in Agile teams, and seven years of independent "
     "contract work requiring consistent output against written guidelines without direct supervision."),
    ("Languages",
     "English, Hindi, Urdu, and Marathi — the basis for multilingual AI evaluation, transcription, "
     "and localization work."),
    ("Portfolio",
     "waseemauto.com — selected projects, full experience, and downloadable CV."),
]

tbl2 = doc.add_table(rows=0, cols=2)
tbl2.alignment = WD_TABLE_ALIGNMENT.LEFT
tbl2.autofit = False
tbl2.columns[0].width = col_w[0]
tbl2.columns[1].width = col_w[1]

for k, v in extras:
    row = tbl2.add_row()
    for i, cell in enumerate(row.cells):
        cell.width = col_w[i]
        cp = cell.paragraphs[0]
        cp.paragraph_format.space_before = Pt(2)
        cp.paragraph_format.space_after = Pt(2)
        cp.paragraph_format.line_spacing = 1.1
    run(row.cells[0].paragraphs[0], k, size=9.2, bold=True)
    run(row.cells[1].paragraphs[0], v, size=9.2, color=MUTED)

tblPr2 = tbl2._tbl.tblPr
borders2 = OxmlElement("w:tblBorders")
for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
    e = OxmlElement(f"w:{edge}")
    e.set(qn("w:val"), "none")
    borders2.append(e)
tblPr2.append(borders2)
ind2 = OxmlElement("w:tblInd")
ind2.set(qn("w:w"), "0")
ind2.set(qn("w:type"), "dxa")
tblPr2.append(ind2)

doc.save("/sessions/funny-keen-archimedes/cv/Waseem-Ansari-CV.docx")
print("saved")
